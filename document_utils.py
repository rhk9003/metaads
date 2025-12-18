from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from datetime import timedelta
import io

def set_run_font(run, size=12, bold=False):
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")

def generate_docx_bytes(party_a, email, payment_opt, start_dt, pay_day, pay_dt, case_num, provider_name="高如慧", bank_name="中國信託商業銀行", bank_code="822", account_number="783540208870"):
    doc = Document()

    # 全文行距
    style = doc.styles["Normal"]
    style.paragraph_format.line_spacing = 1.5

    # 標題
    heading = doc.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = heading.add_run("廣告投放服務合約書")
    set_run_font(run, size=18, bold=True)
    
    # 案件編號
    if case_num:
        sub_head = doc.add_paragraph()
        sub_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sub = sub_head.add_run(f"案件編號：{case_num}")
        set_run_font(run_sub, size=10, bold=False)
    
    doc.add_paragraph("")

    # 變數
    if payment_opt == "17,000元/月（每月付款）":
        end_dt = start_dt + timedelta(days=30)
        period_text = (
            f"自 {start_dt.strftime('%Y 年 %m 月 %d 日')} 起至 {end_dt.strftime('%Y 年 %m 月 %d 日')} 止，共 1 個月。"
            "届期如雙方無異議，則本合約自動續行 1 個月，以此類推。"
        )
        price_text = "1. 甲方同意支付乙方服務費用 新台幣壹萬柒仟元整（NT$17,000）／月。"
        pay_time_text = f"2. 付款時間：甲方應於每月 {pay_day} 日前支付當月服務費用至乙方指定帳戶。"
        first_pay_text = f"3. 首期款項應於合作啟動日（{start_dt.strftime('%Y 年 %m 月 %d 日')}）前支付完成。"
        refund_text = "2. 月付方案：已支付之當期費用不予退還。"
    else:
        end_dt = start_dt + timedelta(days=90)
        period_text = (
            f"自 {start_dt.strftime('%Y 年 %m 月 %d 日')} 起至 {end_dt.strftime('%Y 年 %m 月 %d 日')} 止，共 3 個月。"
            "届期如雙方有意續約，應於届滿前 7 日另行協議。"
        )
        price_text = "1. 甲方同意支付乙方服務費用 新台幣肆萬伍仟元整（NT$45,000）／三個月。"
        pay_time_text = f"2. 付款時間：甲方應於 {pay_dt.strftime('%Y 年 %m 月 %d 日')} 前一次支付完成。"
        first_pay_text = None
        refund_text = (
            "2. 季付方案屬優惠性質之預付服務費，一經支付後即不予退還。"
            "即使甲方於合約期間內提前終止或未使用完畢服務內容，亦同；"
            "惟因乙方重大違約致服務無法履行者，不在此限。"
        )

    # 立約人
    p = doc.add_paragraph()
    run = p.add_run(f"甲方（委託暨付款方）：{party_a}\n")
    set_run_font(run, size=12, bold=True)
    run = p.add_run(f"乙方（服務執行者）：{provider_name}")
    set_run_font(run, size=12, bold=True)
    doc.add_paragraph("")

    # 前言
    p = doc.add_paragraph()
    run = p.add_run("茲因甲方委託乙方提供數位廣告投放服務，雙方本於誠信原則，同意訂立本合約，並共同遵守下列條款：")
    set_run_font(run)

    # 通用條款加入函式
    def add_clause(title, contents):
        p_title = doc.add_paragraph()
        run_title = p_title.add_run(title)
        set_run_font(run_title, size=12, bold=True)

        for content in contents:
            if content:
                p_item = doc.add_paragraph()
                p_item.paragraph_format.left_indent = Cm(0.75)
                run_item = p_item.add_run(content)
                set_run_font(run_item)

    # 第一條
    add_clause("第一條　合約期間", [period_text])

    # 第二條：服務內容
    doc.add_paragraph("")
    p = doc.add_paragraph()
    run = p.add_run("第二條　服務內容")
    set_run_font(run, bold=True)

    p = doc.add_paragraph()
    run = p.add_run("乙方同意為甲方提供以下廣告投放服務：")
    set_run_font(run)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run("一、固定工作項目")
    set_run_font(run, bold=True)

    items_fixed = [
        "1. 廣告上架：依甲方需求於指定平台建立並上架廣告活動。",
        "2. 廣告監控／維護／優化：定期監控成效數據，進行必要之調整與優化。",
        "3. 簡易週報：每週提供廣告成效摘要及下週優化方向。"
    ]
    for item in items_fixed:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1.5)
        set_run_font(p.runs[0])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    run = p.add_run("二、非固定工作項目（視實際狀況提供）")
    set_run_font(run, bold=True)

    items_non = [
        "1. 廣告文案與素材優化：本服務雖以投放操作為主，惟視整體成效需求，乙方得主動提出文案修改建議（如：提供不同版本文案供甲方選擇或修訂）。",
        "2. 網頁調整建議：為確保廣告宣傳訴求一致並協助達成成效，乙方得針對廣告到達頁面（Landing Page）提供調整建議。"
    ]
    for item in items_non:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(1.5)
        set_run_font(p.runs[0])

    # 第三～十四條
    add_clause("第三條　服務範圍與限制", [
        "1. 本服務範圍以 Meta（Facebook／Instagram）廣告投放為主；若需擴展至其他平台，雙方另行協議。",
        "2. 廣告投放預算由甲方自行支付予廣告平台，不包含於本合約服務費用內。",
        "3. 廣告素材（圖片、影片等）之製作原則上由甲方提供，乙方提供方向與建議。",
        "4. 甲方應提供必要帳號權限、素材與資訊，以確保服務得以順利執行。"
    ])

    add_clause("第四條　配合事項與作業方式", [
        "1. 甲方同意配合乙方所需之資料提供、權限設定與必要操作，以確保服務品質。",
        "2. 若因平台政策、帳號狀況或其他不可控因素需採替代作業方式（例如：由甲方匯出報表供乙方監控），甲方同意合理配合。"
    ])

    add_clause("第五條　費用與付款方式", [
        price_text,
        pay_time_text,
        first_pay_text,
        "4. 逾期付款者，乙方得暫停服務至款項付清為止；因此造成之廣告中斷或成效波動，乙方不負賠償責任。"
    ])

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    run = p.add_run(f"乙方指定收款帳戶：\n銀行：{bank_name}（{bank_code}）\n帳號：{account_number}")
    set_run_font(run)

    add_clause("第六條　付款方式與稅務責任", [
        "1. 乙方為自然人，依法無須開立統一發票。",
        "2. 本合約費用之付款方式、帳務處理及相關稅務申報，均由甲方依其自身狀況及相關法令自行決定並負責。",
        "3. 甲方得依其帳務或實務需求，選擇是否以勞務報酬方式支付或其他合法方式付款；乙方將於合理需求下配合提供必要之收款或服務文件。",
        "4. 乙方不負責判斷、建議或保證任何稅務處理方式之合法性。"
    ])

    add_clause("第七條　成效聲明與免責", [
        "1. 乙方將盡專業所能優化廣告成效，但投放成效受市場環境、競爭狀況、消費者行為、平台演算法等多重因素影響，乙方不保證特定之轉換率、ROAS 或銷售成果。",
        "2. 因平台政策變更、帳號異常、不可抗力因素等非乙方可控原因導致之廣告中斷或成效下降，乙方不負賠償責任。",
        "3. 甲方提供之素材、商品或服務如違反平台政策或法令規定，導致廣告被拒絕或帳號受處分，乙方不負相關責任。"
    ])

    add_clause("第八條　保密條款", [
        "1. 合作期間所涉及之商業資訊、廣告數據、行銷策略及客戶資料等，均屬機密資訊，僅得用於本合作目的。",
        "2. 本保密義務於合約終止後仍持續有效 2 年。"
    ])

    add_clause("第九條　智慧財產權", [
        "1. 乙方提供之廣告文案、策略建議、報告等成果，甲方於付清全部款項後，得於本案範圍內使用。",
        "2. 甲方提供之品牌素材、商標、圖片等，其權利仍歸甲方所有。"
    ])

    add_clause("第十條　合約終止", [
        "1. 任一方如欲提前終止本合約，應於終止日前 14 日以書面（含電子郵件、通訊軟體訊息）通知他方。",
        refund_text,
        "3. 如因一方重大違約致他方權益受損，受損方得立即終止合約並請求損害賠償。"
    ])

    add_clause("第十一條　通知方式", [
        "本合約相關通知，得以電子郵件、LINE、Messenger 或其他雙方約定之通訊方式為之，於發送時即生效力。"
    ])

    add_clause("第十二條　合約變更", [
        "本合約之任何修改或補充，應經雙方書面同意後始生效力。"
    ])

    add_clause("第十三條　不可抗力", [
        "因天災、戰爭、政府行為、網路中斷、平台系統異常或其他不可抗力因素，致任一方無法履行本合約義務時，該方不負違約責任；惟應儘速通知並於事由消滅後恢復履行。"
    ])

    add_clause("第十四條　爭議處理", [
        "本合約之解釋與適用，以中華民國法律為準據法。雙方如有爭議，應先行協商；協商不成以臺灣臺北地方法院為第一審管轄法院。"
    ])

    # 簽名欄
    doc.add_paragraph("")
    doc.add_paragraph("")

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False

    cell_a = table.cell(0, 0)
    # 信箱加入簽名欄
    run = cell_a.paragraphs[0].add_run(
        f"甲方（委託暨付款方）：\n{party_a}\n信箱：{email}\n\n簽名：___________________\n\n日期：_____ 年 ___ 月 ___ 日"
    )
    set_run_font(run, size=12)

    cell_b = table.cell(0, 1)
    run = cell_b.paragraphs[0].add_run(
        f"乙方（服務執行者）：\n{provider_name}\n\n簽名：___________________\n\n日期：_____ 年 ___ 月 ___ 日"
    )
    set_run_font(run, size=12)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()
